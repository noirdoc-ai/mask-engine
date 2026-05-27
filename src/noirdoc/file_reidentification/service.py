"""Reidentify pseudonyms in downloaded files.

Replaces <<TYPE_N>> tokens with original values in supported formats:

* **DOCX** – python-docx paragraph runs + table cells
* **XLSX** – openpyxl cell values
* **Plain text** (TXT/CSV/MD/HTML) – simple string replacement

Returns ``None`` for unsupported formats (PDF, PPTX, images) so the
caller can fall through to returning the original bytes.
"""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

import structlog

from noirdoc.mappings.hydration import hydrate_mapper
from noirdoc.reidentification.engine import ReidentificationEngine

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

    from noirdoc.pseudonymization.mapper import PseudonymMapper

logger = structlog.get_logger()

_PSEUDO_PATTERN = re.compile(r"<<[A-Z_]+_\d+>>")

_TEXT_MIMES = {
    "text/plain",
    "text/csv",
    "text/markdown",
    "text/html",
    "text/tab-separated-values",
    "application/json",
}

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def reidentify_file_bytes(
    file_bytes: bytes,
    content_type: str,
    mappings: dict[str, str],
) -> bytes | None:
    """Replace pseudonyms with originals in file bytes.

    Args:
        file_bytes: Raw file content from the provider.
        content_type: MIME type (may include charset parameter).
        mappings: ``{pseudonym: original}`` dict from MappingStore.

    Returns:
        Reidentified bytes, or ``None`` if format is unsupported.
    """
    if not mappings:
        return None

    mapper = hydrate_mapper(mappings)
    engine = ReidentificationEngine()

    # Normalise content_type: strip parameters like "; charset=utf-8"
    mime = content_type.split(";")[0].strip().lower()

    if mime in _TEXT_MIMES:
        return _reidentify_text(file_bytes, engine, mapper)
    if mime == _DOCX_MIME:
        return _reidentify_docx(file_bytes, engine, mapper)
    if mime == _XLSX_MIME:
        return _reidentify_xlsx(file_bytes, engine, mapper)

    return None


def _reidentify_text(
    file_bytes: bytes, engine: ReidentificationEngine, mapper: PseudonymMapper
) -> bytes:
    """Decode → reidentify → encode."""
    text = file_bytes.decode("utf-8", errors="replace")
    reidentified = engine.reidentify(text, mapper)
    return reidentified.encode("utf-8")


def _reidentify_docx(
    file_bytes: bytes, engine: ReidentificationEngine, mapper: PseudonymMapper
) -> bytes | None:
    """Walk DOCX paragraphs and table cells, reidentify text in runs."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        logger.warning("file_reident.docx_load_failed")
        return None

    changed = False

    for para in doc.paragraphs:
        if _reidentify_paragraph(para, engine, mapper):
            changed = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _reidentify_paragraph(para, engine, mapper):
                        changed = True

    if not changed:
        return file_bytes

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _reidentify_paragraph(
    para: Paragraph, engine: ReidentificationEngine, mapper: PseudonymMapper
) -> bool:
    """Reidentify text in a paragraph's runs. Returns True if changed."""
    full_text = para.text
    if not _PSEUDO_PATTERN.search(full_text):
        return False

    new_text = engine.reidentify(full_text, mapper)
    if new_text == full_text:
        return False

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    return True


def _reidentify_xlsx(
    file_bytes: bytes, engine: ReidentificationEngine, mapper: PseudonymMapper
) -> bytes | None:
    """Walk XLSX cells, reidentify string values."""
    from openpyxl import load_workbook

    try:
        wb = load_workbook(io.BytesIO(file_bytes))
    except Exception:
        logger.warning("file_reident.xlsx_load_failed")
        return None

    changed = False

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and _PSEUDO_PATTERN.search(cell.value):
                    new_val = engine.reidentify(cell.value, mapper)
                    if new_val != cell.value:
                        cell.value = new_val
                        changed = True

    if not changed:
        wb.close()
        return file_bytes

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()
