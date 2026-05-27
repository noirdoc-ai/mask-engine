"""DOCX text extraction using python-docx."""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, Protocol

from noirdoc.file_analysis.extractors._zip_safety import check_ooxml_zip_safe

if TYPE_CHECKING:
    from docx.table import Table
    from docx.text.paragraph import Paragraph


class _BlockContainer(Protocol):
    """Common interface of python-docx block containers (body, header, footer, cell)."""

    @property
    def paragraphs(self) -> list[Paragraph]: ...

    @property
    def tables(self) -> list[Table]: ...


def _walk_block_container(container: _BlockContainer, parts: list[str]) -> None:
    """Append non-empty text from every paragraph and table cell in *container*.

    Used for the document body, headers, footers, and comments — all of
    which are BlockItemContainers in python-docx and routinely carry the
    same PII (author lines, "Confidential — Anna Müller" stamps, etc.)
    the body does.
    """
    for para in container.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for cell_para in cell.paragraphs:
                    cell_text = cell_para.text.strip()
                    if cell_text:
                        parts.append(cell_text)


def extract_docx(data: bytes) -> str:
    """Extract text from a DOCX byte-string.

    Walks the document body, all section headers and footers (default,
    first-page, even-page), and review comments. Headers, footers, and
    comments are common PII surfaces that the detector pipeline must see
    before the output is reconstructed.
    """
    from docx import Document

    check_ooxml_zip_safe(data, label="docx")
    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    _walk_block_container(doc, parts)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            _walk_block_container(header, parts)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            _walk_block_container(footer, parts)

    try:
        comments = list(doc.comments)
    except Exception:
        comments = []
    for comment in comments:
        _walk_block_container(comment, parts)

    return "\n".join(parts)
