"""Reconstruct files with pseudonymised content.

Only a subset of formats support in-place reconstruction:

* **DOCX** – replace text in paragraph runs via ``python-docx``
* **XLSX** – replace cell values via ``openpyxl``
* **Plain text** – simple string encode

Formats like PDF and images *cannot* be reconstructed; the caller should
convert those blocks to text instead.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, Protocol

if TYPE_CHECKING:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from noirdoc.file_analysis.models import FileBlock


class _BlockContainer(Protocol):
    """Common interface of python-docx block containers (body, header, footer, cell)."""

    @property
    def paragraphs(self) -> list[Paragraph]: ...

    @property
    def tables(self) -> list[Table]: ...


_RECONSTRUCTABLE_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
    "text/csv",
    "text/markdown",
    "text/html",
}


def can_reconstruct(mime_type: str) -> bool:
    """Return ``True`` if in-place content replacement is supported for *mime_type*."""
    return mime_type in _RECONSTRUCTABLE_MIMES


def reconstruct(block: FileBlock) -> bytes | None:
    """Return new file bytes with pseudonymised text, or ``None`` on failure.

    The caller must ensure ``block.pseudonymized_text`` is set and
    ``can_reconstruct(block.mime_type)`` is ``True``.
    """
    # Short-circuit: pre-computed bytes (e.g. from XLSX column-inference)
    if block.reconstructed_bytes is not None:
        return block.reconstructed_bytes

    if block.pseudonymized_text is None:
        return None

    mime = block.mime_type
    if mime in ("text/plain", "text/csv", "text/markdown", "text/html"):
        return _reconstruct_plain(block)
    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _reconstruct_docx(block)
    if mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return _reconstruct_xlsx(block)
    return None


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


def _reconstruct_plain(block: FileBlock) -> bytes:
    return (block.pseudonymized_text or "").encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _replace_in_block_container(container: _BlockContainer, replacements: dict[str, str]) -> None:
    """Apply *replacements* to every paragraph in *container* and its tables."""
    for para in container.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)


def _reconstruct_docx(block: FileBlock) -> bytes | None:
    """Find-and-replace detected entities in DOCX paragraph runs.

    Covers the document body, every section's headers and footers
    (default, first-page, even-page variants), and review comments —
    matching the surfaces walked by :func:`extract_docx` so PII the
    detector saw is also stripped from the output bytes.
    """
    from docx import Document

    try:
        doc = Document(io.BytesIO(block.content_bytes))
    except Exception:
        return None

    replacements = _build_replacements(block)
    if not replacements:
        return block.content_bytes  # nothing to replace

    _replace_in_block_container(doc, replacements)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            _replace_in_block_container(header, replacements)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            _replace_in_block_container(footer, replacements)

    try:
        comments = list(doc.comments)
    except Exception:
        comments = []
    for comment in comments:
        _replace_in_block_container(comment, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_in_paragraph(para: Paragraph, replacements: dict[str, str]) -> None:
    """Replace entity text across runs in a paragraph."""
    full_text = para.text
    for original, pseudo in replacements.items():
        if original in full_text:
            full_text = full_text.replace(original, pseudo)

    if full_text == para.text:
        return

    # Re-write runs: clear all but first, set first to full text.
    # This simplifies formatting but is acceptable for v1.
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ""


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def _reconstruct_xlsx(block: FileBlock) -> bytes | None:
    """Replace cell values that contain detected entities."""
    from openpyxl import load_workbook

    try:
        wb = load_workbook(io.BytesIO(block.content_bytes))
    except Exception:
        return None

    replacements = _build_replacements(block)
    if not replacements:
        buf = io.BytesIO()
        wb.save(buf)
        wb.close()
        return buf.getvalue()

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    new_val = cell.value
                    for original, pseudo in replacements.items():
                        new_val = new_val.replace(original, pseudo)
                    if new_val != cell.value:
                        cell.value = new_val

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _build_replacements(block: FileBlock) -> dict[str, str]:
    """Map original entity text -> pseudonym using the block's detected entities.

    We rely on the pseudonymized_text already containing the replacements, but
    for DOCX/XLSX reconstruction we need per-entity mapping.  The entities
    store the original ``text`` and the mapper assigned pseudonyms when
    ``pseudonymize()`` was called; we reconstruct the mapping from the entities
    and the pseudonymized_text.
    """
    if not block.entities or not block.extracted_text or not block.pseudonymized_text:
        return {}

    # Simple approach: for each entity, compute what it was replaced with.
    # Entities are sorted by start position; pseudonymize replaces from end to start,
    # so the offsets stay valid.  We rebuild the mapping here.
    replacements: dict[str, str] = {}
    pseudo = block.pseudonymized_text

    # Walk entities and find corresponding pseudonyms via offset tracking
    offset_shift = 0
    for entity in sorted(block.entities, key=lambda e: e.start):
        orig_text = entity.text
        # Find the pseudonym in the pseudonymized text at the shifted position
        pseudo_start = entity.start + offset_shift
        # The pseudonym token looks like <<TYPE_N>> – find it
        if pseudo_start < len(pseudo) and pseudo[pseudo_start:].startswith("<<"):
            end = pseudo.find(">>", pseudo_start)
            if end != -1:
                pseudo_token = pseudo[pseudo_start : end + 2]
                replacements[orig_text] = pseudo_token
                offset_shift += len(pseudo_token) - len(orig_text)

    return replacements
